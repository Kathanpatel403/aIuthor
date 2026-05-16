"""DOCX export from HTML."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)


def html_to_docx(html: str, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from html2docx import html2docx as _html2docx

        doc = _html2docx(html)
        doc.save(str(out_path))
        return
    except ImportError:
        logger.warning("html2docx not installed; using plain-text DOCX fallback")
    except Exception as exc:  # noqa: BLE001
        logger.warning("html2docx failed (%s); using plain-text fallback", exc)
    _plain_text_docx(_html_to_plain(html), out_path)


def _html_to_plain(html: str) -> str:
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "", html)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p>", "\n\n", text)
    text = re.sub(r"(?i)</h[1-6]>", "\n\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _plain_text_docx(text: str, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    for block in text.replace("\r\n", "\n").split("\n\n"):
        for line in block.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph("")
    doc.save(str(out_path))


def markdownish_to_docx(text: str, out_path: Path) -> None:
    """Legacy wrapper for markdown-ish plain text."""
    _plain_text_docx(text, out_path)
